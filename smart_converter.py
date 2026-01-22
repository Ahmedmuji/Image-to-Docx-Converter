import os
import requests
import json
import base64
import docx
from docx import Document
from docx.shared import Inches
import re
import subprocess
import glob
import logging
import streamlit as st
from dotenv import load_dotenv

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Try to get key from Streamlit secrets first (for cloud deployment)
# Fallback to .env (for local development)
try:
    API_KEY = st.secrets["GEMINI_API_KEY"]
except Exception:
    load_dotenv()
    API_KEY = os.getenv("GEMINI_API_KEY")

def get_gemini_response(image_path):
    """
    Sends image to Gemini API and requests text transcription + Python code for diagrams.
    """
    try:
        with open(image_path, "rb") as f:
            image_data = f.read()
            
        base64_image = base64.b64encode(image_data).decode('utf-8')
        
        # Use a stable, specific model version to avoid 'overloaded' errors on experimental/new versions
        model_name = "gemini-1.5-flash"
        logging.info(f"Using model: {model_name}")

        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={API_KEY}"
        headers = {'Content-Type': 'application/json'}
        
        prompt_text = """
        You are an expert OCR and technical diagram transcription system.
        
        TASK:
        1. Transcribe all text from the image accurately.
        2. If you see any diagrams, charts, graphs, or technical illustrations:
           - DO NOT describe them in text.
           - Instead, write a Python script using `matplotlib` to RECREATE that diagram exactly.
           - Place the Python code inside these specific tags: [[DIAGRAM_CODE_START]] ... [[DIAGRAM_CODE_END]]
           - The Python code MUST save the figure to a file named 'generated_diagram.png' and close the plot.
           - Example code structure:
             ```python
             import matplotlib.pyplot as plt
             fig, ax = plt.subplots()
             # ... drawing commands ...
             plt.savefig('generated_diagram.png')
             plt.close()
             ```
           - Use ONLY `matplotlib` and `numpy`.
        
        FORMATTING:
        - Output the text normally.
        - Insert the diagram code blocks in the natural flow where the diagrams appear in the document.
        """
        
        payload = {
            "contents": [{
                "parts": [
                    {"text": prompt_text},
                    {
                        "inline_data": {
                            "mime_type": "image/jpeg", 
                            "data": base64_image
                        }
                    }
                ]
            }]
        }
        
        # Retry logic for 503 errors
        import time
        for attempt in range(3):
            response = requests.post(url, headers=headers, data=json.dumps(payload))
            if response.status_code == 200:
                break
            elif response.status_code == 503:
                logging.warning(f"Model overloaded (503). Retrying attempt {attempt+1}/3...")
                time.sleep(2 * (attempt + 1)) # Exponential backoff
            else:
                logging.error(f"API Error: {response.text}")
                return None
        else:
             logging.error("Failed after 3 retries due to overload.")
             return None

        result = response.json()
        if 'candidates' in result and result['candidates']:
            return result['candidates'][0]['content']['parts'][0]['text']
        else:
            logging.warning("No candidates returned.")
            return None
            
    except Exception as e:
        logging.error(f"Failed to call API: {e}")
        return None

import sys

def execute_diagram_code(code, unique_id):
    """
    Saves and runs the generated python code to produce an image.
    Returns path to the generated image or None.
    """
    # Use relative path so it works on Cloud (Linux) and Local (Windows)
    code_dir = "diagram_code"
    os.makedirs(code_dir, exist_ok=True)
    
    script_path = os.path.join(code_dir, f"diag_{unique_id}.py")
    image_output_name = f"diag_img_{unique_id}.png"
    
    valid_output_path = os.path.join(code_dir, image_output_name)
    # Use forward slashes for compatibility in generated code strings
    escaped_path = valid_output_path.replace('\\', '/')
    
    code = code.replace("generated_diagram.png", escaped_path)
    # Fix common Matplotlib LaTeX errors
    code = code.replace(r"\implies", r"\Rightarrow")
    
    with open(script_path, "w", encoding='utf-8') as f:
        f.write(code)
        
    try:
        logging.info(f"Executing diagram script: {script_path}")
        # Use sys.executable to ensure we use the same python environment (with installed libs)
        result = subprocess.run([sys.executable, script_path], capture_output=True, text=True, timeout=30)
        
        if result.returncode == 0 and os.path.exists(valid_output_path):
            logging.info("Diagram generated successfully.")
            return valid_output_path
        else:
            logging.error(f"Diagram generation failed. Stderr: {result.stderr}")
            return None
    except Exception as e:
        logging.error(f"Error running script: {e}")
        return None



def process_single_image_to_doc(image_path, doc):
    """
    Processes a single image:
    1. Sends to Gemini for text + diagram code
    2. Executes diagram code
    3. Adds content to the provided docx Document object
    """
    filename = os.path.basename(image_path)
    logging.info(f"Processing {filename}...")
    
    response_text = get_gemini_response(image_path)
    if not response_text:
        return False
        
    doc.add_heading(f"Source: {filename}", level=1)
    
    # Parse text and code blocks
    # Split by regex
    parts = re.split(r'\[\[DIAGRAM_CODE_START\]\](.*?)\[\[DIAGRAM_CODE_END\]\]', response_text, flags=re.DOTALL)
    
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # This is TEXT
            if part.strip():
                doc.add_paragraph(part.strip())
        else:
            # This is CODE (odd index)
            code_block = part.strip()
            # Remove markdown fences if present
            code_block = code_block.replace("```python", "").replace("```", "")
            
            unique_id = f"{filename}_{i}"
            # Ensure unique ID is file-system safe
            unique_id = "".join([c for c in unique_id if c.isalpha() or c.isdigit() or c=='_']).rstrip()
            
            img_path = execute_diagram_code(code_block, unique_id)
            
            if img_path:
                try:
                    doc.add_picture(img_path, width=Inches(5))
                except Exception as e:
                    logging.error(f"Could not add image {img_path}: {e}")
                    doc.add_paragraph("[Diagram Generation Failed - See Logs]")
            else:
                doc.add_paragraph("[Diagram Generation Failed]")
    
    doc.add_page_break()
    return True

def main():
    image_dir = "e:/PPIT/test images"
    output_doc_path = os.path.join(image_dir, "Converted_Smart_Output.docx")
    
    doc = Document()
    
    files = [f for f in os.listdir(image_dir) if f.lower().endswith(('.jpg', '.jpeg', '.png'))]
    files.sort()
    
    for filename in files:
        if "WhatsApp" not in filename: continue # Filter for target images if needed
        file_path = os.path.join(image_dir, filename)
        process_single_image_to_doc(file_path, doc)
        doc.save(output_doc_path) # Incremental save
        
    logging.info(f"Done. Saved to {output_doc_path}")

if __name__ == "__main__":
    main()
